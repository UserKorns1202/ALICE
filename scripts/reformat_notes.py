"""
reformat_notes.py

Simple rule-based reformatter to convert verbose markdown notes into a compact
student-note style similar to the provided "Example Notes.md".

Usage:
    python scripts/reformat_notes.py --in "path/to/Leading with Dignity Notes.md" --out out.md

This tool is intentionally lightweight and deterministic: it groups markdown by
headings, keeps the first 1-2 sentences of paragraphs as summaries, preserves
and truncates lists, and emits a small metadata header in the Example Notes
style.
"""
from __future__ import annotations
import argparse
import re
from pathlib import Path
from typing import List, Tuple
try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None
import json
import subprocess
try:
    import requests
except Exception:
    requests = None


def chunk_text(text: str, max_chars: int = 4000, overlap: int = 400):
    start = 0
    n = len(text)
    while start < n:
        end = min(start + max_chars, n)
        yield text[start:end]
        if end == n:
            break
        start = max(0, end - overlap)


def try_ollama_http(prompt: str, model: str = 'llama3.2', max_tokens: int = 2048, temperature: float = 0.3) -> str | None:
    if requests is None:
        return None
    url = 'http://localhost:11434/api/generate'
    payload = {
        'model': model,
        'prompt': prompt,
        'max_tokens': max_tokens,
        'temperature': temperature,
    }
    try:
        # Use streaming to handle Ollama chunked responses
        r = requests.post(url, json=payload, timeout=120, stream=True)
        r.raise_for_status()
        assembled = []
        try:
            for line in r.iter_lines(decode_unicode=True):
                if not line:
                    continue
                line = line.strip()
                try:
                    obj = json.loads(line)
                    # If it's an event object, try common fields
                    if isinstance(obj, dict):
                        for k in ('response', 'text', 'content', 'result'):
                            v = obj.get(k)
                            if isinstance(v, str) and v:
                                assembled.append(v)
                                break
                        # else, try flattening any string values
                        else:
                            for v in obj.values():
                                if isinstance(v, str):
                                    assembled.append(v)
                                    break
                except Exception:
                    # Not JSON — append raw chunk
                    assembled.append(line)
        except Exception:
            pass
        if assembled:
            return '\n'.join(assembled)
        # Fallback to full body text
        try:
            data = r.json()
            if isinstance(data, dict):
                for k in ('response', 'text', 'content', 'result'):
                    if k in data and isinstance(data[k], str):
                        return data[k]
            return r.text
        except Exception:
            return r.text
    except Exception:
        return None


def try_ollama_cli(prompt: str, model: str = 'llama3.2') -> str | None:
    attempts = [
        ['ollama', 'query', model, prompt],
        ['ollama', 'generate', model, '--prompt', prompt],
        ['ollama', 'run', model, '--prompt', prompt],
    ]
    for cmd in attempts:
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            out = proc.stdout.strip()
            if out:
                # Try to parse JSON output if present and extract main text
                try:
                    obj = json.loads(out)
                    if isinstance(obj, dict):
                        for k in ('response', 'text', 'content', 'result'):
                            if k in obj and isinstance(obj[k], str):
                                return obj[k]
                        # fallback: join string values
                        for v in obj.values():
                            if isinstance(v, str):
                                return v
                except Exception:
                    return out
        except Exception:
            continue
    return None


def send_to_ollama(prompt: str, model: str = 'llama3.2', max_tokens: int = 2048, temperature: float = 0.3) -> str:
    out = try_ollama_http(prompt, model=model, max_tokens=max_tokens, temperature=temperature)
    if out:
        return out
    out = try_ollama_cli(prompt, model=model)
    if out:
        return out
    raise RuntimeError('Could not contact Ollama (HTTP or CLI).')


def condense_chunk_prompt(chunk: str, previous_summary: str | None, target_pages: int) -> str:
    prev = f"Previous summary for continuity:\n{previous_summary}\n---\n" if previous_summary else ''
    prompt = (
        f"You are a college senior writing clear, concise, and insightful notes on a book. "
        f"Condense the following excerpt into polished markdown notes suitable for study. "
        f"Output headings and subheadings as necessary. Keep the style reflective, analytical, and precise. "
        f"Aim to contribute to a final document of about {target_pages} pages when combined with other sections: be concise but do not omit essential context.\n"
        f"{prev}Excerpt:\n---\n{chunk}\n---\nProduce the notes now."
    )
    return prompt


def condense_document_with_llm(text: str, model: str = 'llama3.2', target_pages: int = 20) -> str:
    # Chunk the document and summarize each chunk, keeping recent context for continuity
    summaries = []
    prev = ''
    for i, chunk in enumerate(chunk_text(text, max_chars=4000, overlap=400)):
        prompt = condense_chunk_prompt(chunk, prev, target_pages)
        print(f'LLM: processing chunk {i+1}...')
        out = send_to_ollama(prompt, model=model, max_tokens=4096, temperature=0.2)
        out = out.strip()
        summaries.append(out)
        # Keep last 2000 chars for context
        combined = '\n\n'.join(summaries[-3:])
        prev = combined[-2000:]

    # Final refine: ask LLM to combine and condense summaries into final ~target_pages
    combined_all = '\n\n'.join(summaries)
    final_prompt = (
        f"You are to combine the following chunk-level notes into a single coherent, well-organized study document. "
        f"Write in the voice of a college senior analyzing the book: clear, concise, reflective, and insightful. "
        f"Restructure into chapters and subheadings; aim for approximately {target_pages} pages (condense where possible without losing core context). "
        "Preserve key examples and arguments; where repetition occurs, merge and clarify. Output only markdown.\n\n"
        f"Chunk summaries:\n---\n{combined_all}\n---\nFinal condensed notes:" 
    )
    print('LLM: running final refinement...')
    final = send_to_ollama(final_prompt, model=model, max_tokens=8192, temperature=0.2)
    final = final.strip()
    # Ask LLM to perform a cleaning pass: fix broken words, join short-line fragments,
    # and output valid markdown only.
    clean_prompt = (
        "The text below is a markdown document produced by another model. It may contain "
        "broken words (segments split across spaces), excessive single-word lines, and odd line breaks. "
        "Please fix all split words, merge short single-word lines into proper sentences, preserve list items (- ), blockquotes (> ), and headings (lines wrapped in **bold**), and output only the cleaned markdown. "
        "Do not add new content; only repair formatting and minor punctuation. Output only the corrected markdown, nothing else.\n\n"
        f"Document to clean:\n---\n{final}\n---\nCleaned document:"
    )
    try:
        cleaned = send_to_ollama(clean_prompt, model=model, max_tokens=8192, temperature=0.0)
        if cleaned:
            return cleaned.strip()
    except Exception:
        pass

    # Fallback to local cleaning
    return clean_llm_text(final, model_name=model)


_STOPWORDS = {"the","and","in","of","to","a","is","it","for","on","with","as","by","an","be","are","was","were","this","that","i","do"}


def clean_llm_text(text: str, model_name: str | None = None) -> str:
    # Remove model name occurrences at end
    lines = [ln.rstrip() for ln in text.splitlines()]
    # Remove trailing model name if present
    if model_name:
        while lines and lines[-1].strip().lower().startswith(model_name.lower()):
            lines.pop()

    # Merge short consecutive lines into paragraphs (e.g., "Here\nare\nthe\nfinal")
    merged = []
    buffer = []
    for ln in lines:
        if not ln.strip():
            if buffer:
                merged.append(' '.join(buffer))
                buffer = []
            merged.append('')
            continue
        # If line is short (<=12 chars) buffer it, else flush buffer and append line
        if len(ln) <= 12 and ' ' not in ln:
            buffer.append(ln)
        else:
            if buffer:
                merged.append(' '.join(buffer))
                buffer = []
            merged.append(ln)
    if buffer:
        merged.append(' '.join(buffer))

    # Fix chunked words within lines by merging runs of short tokens that are not common stopwords
    out_lines = []
    for ln in merged:
        if not ln:
            out_lines.append('')
            continue
        toks = ln.split()
        i = 0
        new_toks = []
        while i < len(toks):
            # build a segment of consecutive short tokens
            if len(toks[i]) <= 4 and toks[i].lower() not in _STOPWORDS:
                j = i
                seg = [toks[j]]
                j += 1
                while j < len(toks) and len(toks[j]) <= 4 and toks[j].lower() not in _STOPWORDS:
                    seg.append(toks[j])
                    j += 1
                if len(seg) >= 3:
                    # merge segment
                    merged_word = ''.join(seg)
                    new_toks.append(merged_word)
                    i = j
                    continue
            new_toks.append(toks[i])
            i += 1
        out_lines.append(' '.join(new_toks))

    return '\n'.join(out_lines).strip()


SENTENCE_END_RE = re.compile(r"([\.!?])\s+")
HEADING_RE = re.compile(r"^(#{1,6})\s*(.*)$", flags=re.MULTILINE)
LIST_ITEM_RE = re.compile(r"^\s*[-\*+]\s+(.*)$", flags=re.MULTILINE)


def split_sentences(text: str, max_sentences: int = 2) -> str:
    text = text.strip().replace('\n', ' ')
    if not text:
        return ''
    parts = SENTENCE_END_RE.split(text)
    # parts alternates chunk, delimiter, chunk, delimiter ...
    sentences = []
    i = 0
    while i < len(parts) and len(sentences) < max_sentences:
        chunk = parts[i].strip()
        if i + 1 < len(parts):
            delim = parts[i+1]
            sentences.append(chunk + delim)
            i += 2
        else:
            sentences.append(chunk)
            i += 1
    return ' '.join(s for s in sentences if s)


def parse_sections(md: str) -> List[Tuple[str, str]]:
    """Return list of (heading, content) pairs. Top content before first
    heading is returned with heading ''."""
    matches = list(HEADING_RE.finditer(md))
    if not matches:
        return [('', md.strip())]
    sections = []
    prev_end = 0
    # handle leading content
    first = matches[0]
    if first.start() > 0:
        leading = md[0:first.start()].strip()
        sections.append(('', leading))
    for i, m in enumerate(matches):
        level = len(m.group(1))
        title = m.group(2).strip()
        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(md)
        content = md[start:end].strip()
        sections.append((('#' * level) + ' ' + title, content))
    return sections


def extract_list_items(text: str) -> List[str]:
    return [m.group(1).strip() for m in LIST_ITEM_RE.finditer(text)]


def condense_section(heading: str, content: str) -> str:
    out_lines = []
    # Normalize heading style: drop leading #s and use bold-like heading in example
    nice_heading = heading.lstrip('#').strip()
    if nice_heading:
        out_lines.append(f"**{nice_heading}**")
    # Find lists and paragraphs
    items = extract_list_items(content)
    if items:
        # if there are list items, keep up to 6 and summarize long ones
        keep = items[:6]
        for it in keep:
            s = split_sentences(it, max_sentences=2)
            out_lines.append(f"- {s}")
    # For non-list paragraphs, find blocks separated by blank lines
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
    # remove paragraphs that are the same as list items
    paragraphs = [p for p in paragraphs if not LIST_ITEM_RE.search(p)]
    for p in paragraphs[:6]:
        s = split_sentences(p, max_sentences=2)
        if s:
            out_lines.append(s)
    # If nothing captured, add a placeholder
    if not out_lines:
        out_lines.append("- (no concise content)")
    return '\n'.join(out_lines)


def generate_output(sections: List[Tuple[str, str]], title_hint: str) -> str:
    out = []
    out.append("**Name:**\n")
    out.append("**Notes on**\n")
    out.append(f"**{title_hint}**\n")
    out.append("**Preface**\n")
    # If there's leading section content (heading == ''), use first paragraph as preface
    if sections and sections[0][0] == '':
        pre = split_sentences(sections[0][1], max_sentences=3)
        if pre:
            out.append(pre + '\n')

    # Process other sections, but aim for brevity
    for heading, content in sections:
        if heading == '':
            continue
        condensed = condense_section(heading, content)
        out.append(condensed + '\n')

    # final Key Takeaways: gather top 6 strongest lines across sections by length
    all_lines = []
    for _, content in sections:
        for p in re.split(r"\n\s*\n", content):
            s = split_sentences(p, max_sentences=2)
            if s:
                all_lines.append(s)
    # sort by length descending to pick 'major' points
    all_lines = sorted(set(all_lines), key=lambda x: -len(x))
    if all_lines:
        out.append("**Key Takeaways**")
        for ln in all_lines[:6]:
            out.append(f"- {ln}")

    return '\n\n'.join(out)


def write_docx_from_text(text: str, path: Path):
    if Document is None:
        raise RuntimeError("python-docx not installed (pip install python-docx)")
    doc = Document()
    # Basic styling: use default font size for readability
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for line in text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph('')
            continue
        # Bolded headings marked as **Heading**
        if line.startswith('**') and line.endswith('**'):
            h = line.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(h)
            run.bold = True
            continue
        if line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
            continue
        # otherwise plain paragraph
        doc.add_paragraph(line)

    doc.save(str(path))


def main():
    parser = argparse.ArgumentParser(description="Reformat verbose markdown notes into compact Example Notes style")
    parser.add_argument("--in", dest="infile", required=True, help="Input markdown file path")
    parser.add_argument("--out", dest="outfile", help="Output file path (optional)")
    parser.add_argument("--outfile", dest="outfile", help=argparse.SUPPRESS)
    parser.add_argument("--format", choices=("md", "docx"), default="docx", help="Output format (md or docx). Defaults to docx")
    parser.add_argument("--model", default="llama3.2", help="LLM model to use for condensation (Ollama model name)")
    parser.add_argument("--target-pages", type=int, default=20, help="Approximate target page count for condensed output")
    parser.add_argument("--no-llm", dest="use_llm", action='store_false', help="Do not use LLM; use rule-based reformatter instead")
    args = parser.parse_args()

    p = Path(args.infile)
    if not p.exists():
        print("Input file not found:", args.infile)
        return
    md = p.read_text(encoding='utf-8', errors='ignore')
    sections = parse_sections(md)
    title = p.stem
    # If first markdown heading exists, use that as title hint
    for h, _ in sections:
        if h:
            title = h.lstrip('#').strip()
            break
    # If LLM requested, use it to condense the raw markdown into a compact document
    if args.use_llm:
        try:
            condensed = condense_document_with_llm(md, model=args.model, target_pages=args.target_pages)
            condensed = clean_llm_text(condensed, model_name=args.model)
        except Exception as e:
            print("LLM condensation failed:", e)
            print("Falling back to rule-based reformatter.")
            condensed = generate_output(sections, title)
    else:
        condensed = generate_output(sections, title)

    if args.format == 'docx':
        out_path = Path(args.outfile) if args.outfile else p.parent / f"{p.stem}_reformatted.docx"
        write_docx_from_text(condensed, out_path)
    else:
        out_path = Path(args.outfile) if args.outfile else p.parent / f"{p.stem}_reformatted.md"
        out_path.write_text(condensed, encoding='utf-8')

    print("Wrote reformatted notes to:", out_path)


if __name__ == '__main__':
    main()
